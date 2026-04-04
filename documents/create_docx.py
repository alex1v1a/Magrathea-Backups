from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def main():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add logo
    try:
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_para.add_run()
        run.add_picture('documents/vectarr-logo.png', width=Inches(1.5))
    except:
        pass
    
    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('VECTARR LLC')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'
    
    # Company info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('A Texas Limited Liability Company\nvectarr.com')
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    # Horizontal line (border)
    p = doc.add_paragraph()
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    doc.add_paragraph()  # Spacing
    
    # Date
    p = doc.add_paragraph()
    run = p.add_run('Date: ___________________')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # RE line
    p = doc.add_paragraph()
    run = p.add_run('RE: Ownership and Authorization Confirmation — Vehicle Registration\nMember: Marc Alexander Sferrazza')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Content
    p = doc.add_paragraph()
    run = p.add_run('To Whom It May Concern:')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    p = doc.add_paragraph()
    run = p.add_run('This letter serves as official confirmation of the ownership structure and authorization status of ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = p.add_run('Vectarr LLC')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = p.add_run(', a Texas Limited Liability Company (the "Company").')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Ownership Structure
    p = doc.add_paragraph()
    run = p.add_run('OWNERSHIP STRUCTURE')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run('Vectarr LLC is equally owned by two (2) members, each holding a fifty percent (50%) ownership interest in the Company:')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    # List members
    members = [
        'Marc Alexander Sferrazza — 50% Member Interest',
        'Kamal Katul — 50% Member Interest'
    ]
    
    for i, member in enumerate(members, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(member)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Authorization
    p = doc.add_paragraph()
    run = p.add_run('AUTHORIZATION')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run('Marc Alexander Sferrazza is hereby confirmed as an authorized signer and agent of Vectarr LLC, with full authority to:')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    auth_items = [
        'Enter into contracts and agreements on behalf of the Company;',
        'Open, maintain, and operate business banking and financial accounts;',
        'Register, title, and insure motor vehicles in the Company\'s name;',
        'Execute all documentation related to vehicle acquisition, registration, and ownership;',
        'Represent the Company in business transactions and negotiations;',
        'Execute documents, instruments, and obligations binding the Company;',
        'Act as the Company\'s representative in all official capacities.'
    ]
    
    for item in auth_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    p = doc.add_paragraph()
    run = p.add_run('This authorization is made pursuant to the Operating Agreement of Vectarr LLC and the applicable laws of the State of Texas governing limited liability companies.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    p = doc.add_paragraph()
    run = p.add_run('This letter is executed by the undersigned members of Vectarr LLC to confirm the foregoing statements and to provide any third party with reasonable assurance of Marc Alexander Sferrazza\'s authority to act on behalf of the Company.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    p = doc.add_paragraph()
    run = p.add_run('Sincerely,')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing
    
    # Signature table
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # First row - signature lines
    cell1 = table.cell(0, 0)
    cell2 = table.cell(0, 1)
    
    # Add horizontal line border to top of second row cells
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p1.add_run('_' * 40)
    run.font.name = 'Times New Roman'
    
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run('_' * 40)
    run.font.name = 'Times New Roman'
    
    # Second row - names and titles
    cell3 = table.cell(1, 0)
    cell4 = table.cell(1, 1)
    
    p3 = cell3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p3.add_run('Kamal Katul')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = p3.add_run('\nMember, Vectarr LLC\nDate: _______________')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p4 = cell4.paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p4.add_run('Marc Alexander Sferrazza')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = p4.add_run('\nMember, Vectarr LLC\nDate: _______________')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(3.5)
        row.cells[1].width = Inches(3.5)
    
    # Save document
    doc.save('documents/Vectarr-Authorization-Letter.docx')
    print('Document saved successfully!')

if __name__ == '__main__':
    main()
