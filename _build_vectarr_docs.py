from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

logo_path = r"C:\Users\Admin\.openclaw\media\inbound\ec3c296e-de0e-421f-b466-6daad2b57fca.png"
address = "5900 Balcones Drive STE 100 Austin, TX 78731 USA"

content = {
    r"C:\Users\Admin\.openclaw\workspace\vectarr_document_1.docx": [
        ("SHORT-FORM FREELANCER NON-DISCLOSURE AGREEMENT", True),
        ("(Vectarr LLC)", False),
        ("This Agreement is entered into as of [Effective Date] by and between Vectarr LLC (\u201cCompany\u201d) and [Freelancer Name] (\u201cFreelancer\u201d).", False),
        ("1. Confidential Information", True),
        ("Freelancer may receive confidential, proprietary, or export-controlled information belonging to Company or its clients (\u201cConfidential Information\u201d). All such information is confidential whether or not marked.", False),
        ("2. Use and Non-Disclosure", True),
        ("Freelancer shall use Confidential Information solely to perform services for Company and shall not disclose it to any third party.", False),
        ("3. Export Control / ITAR", True),
        ("Freelancer acknowledges that certain information may be subject to U.S. export control laws, including ITAR and EAR. Freelancer certifies they are a U.S. Person unless otherwise disclosed in writing and agrees not to disclose controlled information to any foreign person.", False),
        ("4. Ownership", True),
        ("All work product created by Freelancer is a work made for hire and owned exclusively by Company.", False),
        ("5. Term", True),
        ("Confidentiality obligations survive for five (5) years or as long as required by law.", False),
        ("6. Remedies", True),
        ("Company is entitled to injunctive relief for breach.", False),
        ("Governing Law: As specified in the applicable client agreement.", False),
        ("", False),
        ("Agreed:", True),
        ("Vectarr LLC __________________ Date: ____", False),
        ("Freelancer ___________________ Date: ____", False),
    ],
    r"C:\Users\Admin\.openclaw\workspace\vectarr_document_2.docx": [
        ("BACK-TO-BACK NON-DISCLOSURE AGREEMENT", True),
        ("(Subcontractor Confidentiality Agreement)", False),
        ("This Non-Disclosure Agreement (\u201cAgreement\u201d) is entered into as of [Effective Date], by and between:", False),
        ("[Your Company Legal Name], a [State/Country] [Entity Type], with its principal place of business at [Address] (\u201cCompany\u201d)", False),
        ("and", False),
        ("[Subcontractor Legal Name], a [State/Country] [Entity Type / Individual], with its principal place of business at [Address] (\u201cSubcontractor\u201d).", False),
        ("1. Purpose", True),
        ("Company has entered into one or more agreements with its client(s) (\u201cClient\u201d) that include confidentiality, non-disclosure, data protection, and/or intellectual property obligations (\u201cClient NDA\u201d). Subcontractor is being engaged to perform services for Company that may require access to Confidential Information solely to fulfill Company\u2019s obligations to Client.", False),
        ("2. Definition of Confidential Information", True),
        ("\u201cConfidential Information\u201d means all information, whether disclosed orally, visually, electronically, or in writing, including but not limited to client information, trade secrets, source code, business plans, pricing, financial data, technical information, and personal or regulated data. Confidential Information includes information disclosed by Client directly or indirectly through Company.", False),
        ("3. Flow-Down of Client NDA", True),
        ("All confidentiality obligations owed by Company to Client apply equally to Subcontractor. Where obligations conflict, the most restrictive obligation governs.", False),
        ("4. Use and Non-Disclosure Obligations", True),
        ("Subcontractor shall use Confidential Information solely for the authorized purpose and shall not disclose it to any third party.", False),
        ("5. No Further Subcontracting", True),
        ("Subcontractor may not subcontract or delegate without prior written consent from Company.", False),
        ("6. Security Requirements", True),
        ("Subcontractor shall implement administrative, technical, and physical safeguards consistent with industry best practices.", False),
        ("7. Breach Notification", True),
        ("Subcontractor shall notify Company within 24 hours of any actual or suspected breach.", False),
        ("8. Ownership and Intellectual Property", True),
        ("All Confidential Information remains the exclusive property of Client or Company.", False),
        ("9. Term and Survival", True),
        ("Confidentiality obligations survive termination for the longest period required by law or Client NDA.", False),
        ("10. Governing Law", True),
        ("This Agreement shall be governed by the laws specified in the Client NDA.", False),
        ("", False),
        ("IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.", True),
    ],
    r"C:\Users\Admin\.openclaw\workspace\vectarr_document_3.docx": [
        ("MASTER SUBCONTRACTOR AGREEMENT", True),
        ("(Vectarr LLC)", False),
        ("This Master Subcontractor Agreement (\u201cAgreement\u201d) is entered into as of [Effective Date] by and between Vectarr LLC (\u201cCompany\u201d) and [Subcontractor Name] (\u201cSubcontractor\u201d).", False),
        ("1. Scope of Work", True),
        ("Subcontractor shall perform services as defined in one or more written Statements of Work (\u201cSOW\u201d).", False),
        ("2. Confidentiality", True),
        ("Subcontractor shall protect all confidential and client information and comply with all flowed-down NDA obligations.", False),
        ("3. Export Control / ITAR", True),
        ("Subcontractor agrees to comply with ITAR, EAR, and all applicable export control laws. ITAR-controlled data shall only be accessed by authorized U.S. Persons.", False),
        ("4. Intellectual Property", True),
        ("All deliverables, inventions, and work product are works made for hire and owned exclusively by Company.", False),
        ("5. No Subcontracting", True),
        ("Subcontractor may not further subcontract without prior written consent.", False),
        ("6. Security", True),
        ("Subcontractor shall maintain industry-standard administrative, technical, and physical safeguards.", False),
        ("7. Indemnification", True),
        ("Subcontractor shall indemnify and hold harmless Company and its clients from all claims arising from breach.", False),
        ("8. Term and Termination", True),
        ("This Agreement remains in effect until terminated by either party upon written notice. Obligations survive termination.", False),
        ("9. Governing Law", True),
        ("Governing law shall match the applicable client agreement.", False),
        ("", False),
        ("IN WITNESS WHEREOF, the parties agree to this Agreement.", True),
        ("Vectarr LLC __________________ Date: ____", False),
        ("Subcontractor ________________ Date: ____", False),
    ],
}


def set_run_font(run, size=11, bold=False, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)


def create_doc(path, lines):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = ""

    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = True

    cell_logo = table.cell(0, 0)
    p_logo = cell_logo.paragraphs[0]
    run_logo = p_logo.add_run()
    run_logo.add_picture(logo_path, width=Inches(1.4))
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cell_addr = table.cell(0, 1)
    p_addr = cell_addr.paragraphs[0]
    run_addr = p_addr.add_run(address)
    set_run_font(run_addr, size=9, bold=False)
    p_addr.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Space after header
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)

    for text, is_header in lines:
        if text == "":
            doc.add_paragraph()
            continue
        para = doc.add_paragraph()
        run = para.add_run(text)
        if is_header:
            set_run_font(run, size=13, bold=True)
        else:
            set_run_font(run, size=11, bold=False)
        para.paragraph_format.space_after = Pt(6)

    doc.save(path)


for path, lines in content.items():
    create_doc(path, lines)

print("Documents created.")
